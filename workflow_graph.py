from langgraph import Graph, Node
from agents_tools import planner_tool, search_tool, writer_tool
from docx import Document
from datetime import datetime

planner_node = Node(
    func=planner_tool.run,
    name="Planner",
    description="Generates research questions from a topic."
)

search_node = Node(
    func=search_tool.run,
    name="Searcher",
    description="Fetches web content for each research question."
)

writer_node = Node(
    func=writer_tool.run,
    name="Writer",
    description="Generates detailed report with citations."
)

research_graph = Graph()

research_graph.add_edge(planner_node, search_node)
research_graph.add_edge(search_node, writer_node)

def run_research_workflow(topic: str):

    questions_text = planner_node.run(topic)
    questions = [q.strip() for q in questions_text.split("\n") if q.strip() and q[0].isdigit()]

    search_results = []
    for q in questions:
        search_results.append(search_node.run(q))

    report = writer_node.run(topic, search_results, batch_size=3)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Research_Report_{timestamp}.docx"
    doc = Document()
    doc.add_heading(f"Research Report: {topic}", level=0)
    doc.add_paragraph(report)
    doc.save(filename)

    print(f"\nResearch report saved as '{filename}'")
    return report

if __name__ == "__main__":
    topic_input = input("Enter your research topic: ")
    final_report = run_research_workflow(topic_input)
    print("\n==================== FINAL PROFESSIONAL RESEARCH REPORT ====================\n")
    print(final_report)

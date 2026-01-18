from planner import planner_agent
from search_agent import search_agent
from writer_agent import writer_agent
from docx import Document
from datetime import datetime

def generate_report(topic: str):
    print(f"\n[STEP 1] Generating research questions for topic: {topic}\n")
    questions_text = planner_agent(topic)

    questions = [q.strip() for q in questions_text.split("\n") if q.strip() and q[0].isdigit()]
    print(f"Generated {len(questions)} research questions.\n")

    search_results = []
    print("[STEP 2] Retrieving content from web for each question...\n")
    for idx, q in enumerate(questions, start=1):
        print(f"Fetching content for Q{idx}: {q}")
        result = search_agent(q)
        search_results.append(result)

    print("\n[STEP 3] Generating detailed research report with in-text citations...\n")
    report_content = writer_agent(topic, search_results, batch_size=3)

    # Save report as Word document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Research_Report_{timestamp}.docx"
    doc = Document()
    doc.add_heading(f"Research Report: {topic}", level=0)
    doc.add_paragraph(report_content)
    doc.save(filename)

    print(f"\nReport saved successfully as '{filename}'\n")
    return report_content

if __name__ == "__main__":
    topic_input = input("Enter your research topic: ")
    final_report = generate_report(topic_input)
    print("\n==================== FINAL PROFESSIONAL RESEARCH REPORT ====================\n")
    print(final_report)
